"""
Módulo de Manipulação de Documentos DOCX para Mail Merge
Permite carregar, visualizar, editar e personalizar documentos Word com dados de planilhas.
"""

import re
import io
import base64
from pathlib import Path
from typing import Optional, Dict, List, Tuple, Any
from dataclasses import dataclass, field
import tempfile
import zipfile

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None

try:
    import mammoth
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False


@dataclass
class DocxField:
    """Representa um campo placeholder no documento DOCX."""
    name: str
    occurrences: int = 1
    paragraph_indices: List[int] = field(default_factory=list)
    run_indices: List[int] = field(default_factory=list)


@dataclass
class DocxTemplate:
    """Template de documento DOCX com placeholders."""
    file_path: Optional[Path] = None
    document: Optional[Any] = None
    fields: List[DocxField] = field(default_factory=list)
    html_content: str = ""
    raw_text: str = ""
    images: Dict[str, bytes] = field(default_factory=dict)
    is_loaded: bool = False


class DocxDataSource:
    """Gerencia carregamento e manipulação de documentos DOCX."""
    
    FIELD_PATTERN = re.compile(r"{{\s*([^{}]+?)\s*}}")
    
    def __init__(self):
        self.template: Optional[DocxTemplate] = None
        self.file_path: Optional[Path] = None
        
    def load_document(self, file_path: str) -> Tuple[bool, str]:
        """
        Carrega um documento DOCX e retorna (sucesso, mensagem).
        """
        if not HAS_DOCX:
            return False, "python-docx não está instalado. Execute: pip install python-docx"
        
        try:
            path = Path(file_path)
            if path.suffix.lower() != ".docx":
                return False, "Arquivo não é um documento Word (.docx)"
            
            if not path.exists():
                return False, f"Arquivo não encontrado: {file_path}"
            
            # Carregar documento
            doc = Document(str(path))
            
            # Extrair texto completo
            raw_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extrair campos placeholder
            fields = self._extract_fields(doc)
            
            # Converter para HTML para visualização
            html_content = self._convert_to_html(str(path))
            
            # Extrair imagens
            images = self._extract_images(path)
            
            self.template = DocxTemplate(
                file_path=path,
                document=doc,
                fields=fields,
                html_content=html_content,
                raw_text=raw_text,
                images=images,
                is_loaded=True
            )
            
            self.file_path = path
            
            return True, f"Documento carregado: {len(fields)} campos encontrados"
            
        except zipfile.BadZipFile:
            return False, "Arquivo DOCX corrompido ou inválido"
        except Exception as ex:
            return False, f"Erro ao carregar documento: {str(ex)}"
    
    def _extract_fields(self, doc: Any) -> List[DocxField]:
        """Extrai todos os campos {{...}} do documento."""
        field_map: Dict[str, DocxField] = {}
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            matches = self.FIELD_PATTERN.findall(text)
            
            for field_name in matches:
                if field_name not in field_map:
                    field_map[field_name] = DocxField(name=field_name)
                field_map[field_name].occurrences += 1
                field_map[field_name].paragraph_indices.append(para_idx)
        
        return list(field_map.values())
    
    def _convert_to_html(self, file_path: str) -> str:
        """Converte DOCX para HTML usando mammoth."""
        if not HAS_MAMMOTH:
            # Fallback: converte texto simples para HTML básico
            if self.template and self.template.raw_text:
                paragraphs = self.template.raw_text.split("\n")
                html_parts = [f"<p>{p}</p>" for p in paragraphs if p.strip()]
                return "\n".join(html_parts)
            return "<p>Conteúdo não disponível</p>"
        
        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                return result.value
        except Exception:
            return "<p>Erro na conversão para HTML</p>"
    
    def _extract_images(self, path: Path) -> Dict[str, bytes]:
        """Extrai imagens incorporadas no DOCX."""
        images = {}
        try:
            with zipfile.ZipFile(path, 'r') as zip_ref:
                # Imagens estão em word/media/
                for name in zip_ref.namelist():
                    if name.startswith("word/media/"):
                        images[name] = zip_ref.read(name)
        except Exception:
            pass
        return images
    
    def get_preview_html(self, context: Optional[Dict[str, str]] = None) -> str:
        """
        Retorna HTML do documento com campos substituídos (se context fornecido).
        """
        if not self.template or not self.template.is_loaded:
            return "<p>Nenhum documento carregado</p>"
        
        html = self.template.html_content
        
        if context:
            # Substituir campos pelos valores do contexto
            for field_name, value in context.items():
                placeholder = f"{{{{{field_name}}}}}"
                html = html.replace(placeholder, str(value))
        
        return html
    
    def render_document(self, row_data: Dict[str, str]) -> Tuple[str, str]:
        """
        Renderiza o documento com dados de uma linha.
        Retorna (texto_renderizado, html_renderizado).
        """
        if not self.template or not self.template.document:
            return "", ""
        
        # Criar cópia do documento para modificar
        from io import BytesIO
        output = BytesIO()
        
        # Renderizar texto
        rendered_text = self.template.raw_text
        rendered_html = self.template.html_content
        
        for field_name, value in row_data.items():
            placeholder = f"{{{{{field_name}}}}}"
            rendered_text = rendered_text.replace(placeholder, str(value))
            rendered_html = rendered_html.replace(placeholder, str(value))
        
        # Salvar documento modificado
        try:
            # Criar novo documento com texto renderizado
            new_doc = Document()
            
            # Copiar estrutura do original
            for paragraph in self.template.document.paragraphs:
                new_para = new_doc.add_paragraph()
                new_para.style = paragraph.style
                
                # Processar runs (partes formatadas do parágrafo)
                for run in paragraph.runs:
                    new_run = new_para.add_run()
                    new_run.text = run.text
                    
                    # Copiar formatação
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    
                    # Substituir placeholders no run
                    for field_name, value in row_data.items():
                        placeholder = f"{{{{{field_name}}}}}"
                        if placeholder in new_run.text:
                            new_run.text = new_run.text.replace(placeholder, str(value))
            
            new_doc.save(output)
            output.seek(0)
            
            # Converter para base64 para anexar no email
            doc_base64 = base64.b64encode(output.read()).decode('utf-8')
            
            return rendered_text, doc_base64
            
        except Exception as ex:
            return rendered_text, ""
    
    def get_fields_list(self) -> List[str]:
        """Retorna lista de campos disponíveis no documento."""
        if not self.template:
            return []
        return [field.name for field in self.template.fields]
    
    def validate_fields(self, available_columns: List[str]) -> Tuple[bool, List[str]]:
        """
        Valida se todos os campos do documento existem nas colunas da planilha.
        Retorna (válido, lista_de_campos_faltantes).
        """
        if not self.template:
            return False, ["Nenhum documento carregado"]
        
        missing_fields = []
        doc_fields = self.get_fields_list()
        
        for field_name in doc_fields:
            if field_name not in available_columns:
                missing_fields.append(field_name)
        
        return len(missing_fields) == 0, missing_fields
    
    def insert_field_at_cursor(self, current_text: str, field_name: str) -> str:
        """Insere um placeholder de campo no texto."""
        placeholder = f"{{{{{field_name}}}}}"
        separator = "" if not current_text or current_text.endswith((" ", "\n")) else " "
        return f"{current_text}{separator}{placeholder}"
    
    def get_document_info(self) -> Dict[str, Any]:
        """Retorna informações sobre o documento carregado."""
        if not self.template or not self.template.is_loaded:
            return {}
        
        return {
            "file_name": self.template.file_path.name if self.template.file_path else "",
            "fields_count": len(self.template.fields),
            "fields": [field.name for field in self.template.fields],
            "paragraphs_count": len(self.template.document.paragraphs) if self.template.document else 0,
            "has_images": len(self.template.images) > 0,
            "file_size": self.template.file_path.stat().st_size if self.template.file_path else 0,
        }


class DocxMergeEngine:
    """Motor de processamento de mail merge com documentos DOCX."""
    
    def __init__(self):
        self.data_source = DocxDataSource()
        self.processed_documents: List[Dict] = []
    
    def load_template(self, file_path: str) -> Tuple[bool, str]:
        """Carrega um template DOCX."""
        success, message = self.data_source.load_document(file_path)
        return success, message
    
    def get_preview_html(self, context: Optional[Dict[str, str]] = None) -> str:
        """Retorna preview HTML do documento."""
        return self.data_source.get_preview_html(context)
    
    def process_batch(self, row_iterator, max_records: Optional[int] = None) -> Tuple[int, List[str]]:
        """
        Processa múltiplas linhas gerando documentos personalizados.
        Retorna (total_processado, lista_de_erros).
        """
        if not self.data_source.template or not self.data_source.template.is_loaded:
            return 0, ["Template DOCX não carregado"]
        
        self.processed_documents = []
        errors = []
        processed = 0
        
        try:
            for i, row_data in enumerate(row_iterator):
                if max_records and i >= max_records:
                    break
                
                try:
                    rendered_text, doc_base64 = self.data_source.render_document(row_data)
                    
                    self.processed_documents.append({
                        "row_index": i,
                        "rendered_text": rendered_text,
                        "doc_base64": doc_base64,
                        "row_data": row_data,
                    })
                    processed += 1
                except Exception as ex:
                    errors.append(f"Linha {i+1}: {str(ex)}")
        
        except Exception as ex:
            errors.append(f"Erro ao processar batch: {str(ex)}")
        
        return processed, errors
    
    def get_document_info(self) -> Dict[str, Any]:
        """Retorna informações do documento."""
        return self.data_source.get_document_info()
    
    def validate_with_columns(self, columns: List[str]) -> Tuple[bool, List[str]]:
        """Valida campos do documento com colunas disponíveis."""
        return self.data_source.validate_fields(columns)


def create_docx_with_placeholders(template_path: str, output_path: str, replacements: Dict[str, str]) -> bool:
    """
    Função utilitária para criar um DOCX com substituições.
    """
    if not HAS_DOCX:
        return False
    
    try:
        # Carregar template
        engine = DocxMergeEngine()
        success, _ = engine.load_template(template_path)
        
        if not success:
            return False
        
        # Renderizar com dados
        rendered_text, doc_base64 = engine.data_source.render_document(replacements)
        
        # Salvar arquivo
        if doc_base64:
            doc_bytes = base64.b64decode(doc_base64)
            with open(output_path, 'wb') as f:
                f.write(doc_bytes)
            return True
        
        return False
        
    except Exception:
        return False